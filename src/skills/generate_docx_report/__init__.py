import asyncio
import json
from pathlib import Path

from docx import Document

from src.core.runtime_context import get_prompt_config


def _build_document(output_path: Path, title: str, content: str) -> None:
    document = Document()
    document.add_heading(title, level=0)

    paragraphs = [part.strip() for part in str(content).split("\n\n") if part.strip()]
    if not paragraphs:
        paragraphs = [str(content).strip()]

    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    document.save(str(output_path))


def _resolve_downloads_dir() -> Path:
    prompt_config = get_prompt_config()
    raw_downloads_dir = getattr(prompt_config, "downloads_dir", None)
    if not raw_downloads_dir:
        raise RuntimeError(
            "Prompt configuration is unavailable; cannot resolve downloads directory"
        )
    return Path(raw_downloads_dir).expanduser().resolve()


async def run(title: str, content: str, filename: str) -> str:
    try:
        if not isinstance(filename, str) or not filename.strip():
            return json.dumps({
                "status": "error",
                "message": "filename must be a non-empty string",
            })

        requested_name = filename.strip()
        if not requested_name.lower().endswith(".docx"):
            return json.dumps({
                "status": "error",
                "message": "filename must end with .docx",
            })

        normalized_name = Path(requested_name).name
        if normalized_name != requested_name:
            return json.dumps({
                "status": "error",
                "message": "filename must not include directory components",
            })

        downloads_dir = _resolve_downloads_dir()
        downloads_dir.mkdir(parents=True, exist_ok=True)

        output_path = (downloads_dir / normalized_name).resolve()
        try:
            output_path.relative_to(downloads_dir)
        except ValueError:
            return json.dumps({
                "status": "error",
                "message": "Output path is outside configured downloads directory",
            })

        await asyncio.to_thread(
            _build_document,
            output_path,
            str(title),
            str(content),
        )

        return json.dumps({
            "status": "success",
            "file_path": str(output_path),
        })
    except asyncio.CancelledError:
        raise
    except Exception as exc:
        return json.dumps({
            "status": "error",
            "message": f"Failed to generate document: {exc}",
        })


async def generate_docx_report(title: str, content: str, filename: str) -> str:
    return await run(title=title, content=content, filename=filename)
